import requests
import textract
from bs4 import BeautifulSoup
from docx import Document


def parse_url(link):
    html = requests.get(link).text
    soup = BeautifulSoup(html, "html.parser")
    res = soup.get_text()
    return res


def parse_file(filename):
    text = textract.process(filename)
    return text.decode('utf-8')


def save_file(filename, stats, common):
    document = Document()

    document.add_heading('Stats', level=2)
    rows_stats = len(stats)
    cols_stats = len(stats[0])
    table_stats = document.add_table(rows=rows_stats, cols=cols_stats, style="Table Grid")
    for i in range(rows_stats):
        for j in range(cols_stats):
            table_stats.rows[i].cells[j].text = str(stats[i][j])

    document.add_heading('Top 50 common words', level=2)
    rows_common = len(common)
    cols_common = 2
    table_common = document.add_table(rows=rows_common, cols=cols_common, style="Table Grid")
    row = 0
    for key, value in common:
        table_common.rows[row].cells[0].text = str(key)
        table_common.rows[row].cells[1].text = str(value)
        row = row + 1

    document.save(filename)
